# Processing engine for FormMonkey document analysis
# Uses Redis (or in-memory fallback) for job state persistence

import os
import asyncio
import fitz  # PyMuPDF
import docx
from typing import Dict, Any, Optional, Union
from datetime import datetime
import json
import traceback

from shared.types import JobData, ProcessingStatus
from storage.job_store import store_job, get_job, update_job, update_job_status

async def process_document(file_path: str, job_id: str, user_id: str) -> JobData:
    """
    Asynchronously process a document, extracting text from PDF (using PyMuPDF) or DOCX (using python-docx).
    
    Args:
        file_path: Path to the uploaded document
        job_id: Unique identifier for this processing job
        user_id: ID of the user who uploaded the document
        
    Returns:
        JobData containing processing status and extracted content
    """
    # Initialize job in job store
    job_data = JobData(
        job_id=job_id,
        status=ProcessingStatus.PENDING,
        progress=0,
        user_id=user_id,
        file_path=file_path,
        content="",
        error=None,
        created_at=datetime.now().isoformat(),
        updated_at=datetime.now().isoformat()
    )
    await store_job(job_data)
    
    try:
        # Update status to processing
        await update_job_status(job_id, ProcessingStatus.PROCESSING, progress=10)
        
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # Extract text based on file type
        content = ""
        if ext == ".pdf":
            content = await extract_text_from_pdf(file_path, job_id)
        elif ext == ".docx":
            content = await extract_text_from_docx(file_path, job_id)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
        
        # Update content in job store
        await update_job(job_id, {"content": content})
            
        # Mark as completed
        await update_job_status(job_id, ProcessingStatus.COMPLETED, progress=100)
        
    except Exception as e:
        # Handle errors with detailed information
        error_detail = f"{str(e)}\n{traceback.format_exc()}"
        await update_job_status(
            job_id, 
            ProcessingStatus.FAILED, 
            progress=0,
            error=error_detail
        )
        raise
    
    # Return the updated job data
    return await get_job(job_id)

async def extract_text_from_pdf(file_path: str, job_id: str) -> str:
    """Extract text from a PDF file using PyMuPDF"""
    await update_job_status(job_id, ProcessingStatus.PROCESSING, progress=20)
    
    text = ""
    try:
        # Process in a separate thread to avoid blocking
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, _extract_pdf_text, file_path, job_id)
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")
    
    await update_job_status(job_id, ProcessingStatus.PROCESSING, progress=80)
    return text

def _extract_pdf_text(file_path: str, job_id: str) -> str:
    """Non-async PDF text extraction to run in executor"""
    text = ""
    progress_update_task = None
    
    try:
        doc = fitz.open(file_path)
        total_pages = len(doc)
        
        # Create a function to update progress asynchronously
        async def update_progress(page_num, total):
            progress = 20 + int(60 * (page_num + 1) / total)
            await update_job_status(job_id, ProcessingStatus.PROCESSING, progress=progress)
        
        for i, page in enumerate(doc):
            text += page.get_text()
            
            # Update progress every 5 pages or on first/last page
            if i == 0 or i == total_pages - 1 or i % 5 == 0:
                # Create a new event loop for this thread if needed
                try:
                    loop = asyncio.get_event_loop()
                except RuntimeError:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                
                # Wait for any previous update to complete
                if progress_update_task is not None:
                    if not progress_update_task.done():
                        loop.run_until_complete(progress_update_task)
                
                # Schedule the new progress update
                progress_update_task = loop.create_task(update_progress(i, total_pages))
    
    except Exception as e:
        raise Exception(f"PDF extraction error: {str(e)}")
    finally:
        if 'doc' in locals():
            doc.close()
    
    return text

async def extract_text_from_docx(file_path: str, job_id: str) -> str:
    """Extract text from a DOCX file using python-docx"""
    await update_job_status(job_id, ProcessingStatus.PROCESSING, progress=20)
    
    text = ""
    try:
        # Process in a separate thread to avoid blocking
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, _extract_docx_text, file_path, job_id)
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")
    
    await update_job_status(job_id, ProcessingStatus.PROCESSING, progress=80)
    return text

def _extract_docx_text(file_path: str, job_id: str) -> str:
    """Non-async DOCX text extraction to run in executor"""
    text = ""
    try:
        doc = docx.Document(file_path)
        
        # Count paragraphs and tables for progress tracking
        total_elements = len(doc.paragraphs) + len(doc.tables)
        processed = 0
        progress_update_task = None
        
        # Create a function to update progress asynchronously
        async def update_progress(current, total):
            progress = 20 + int(60 * current / total) if total > 0 else 50
            await update_job_status(job_id, ProcessingStatus.PROCESSING, progress=progress)
        
        # Process paragraphs
        for i, para in enumerate(doc.paragraphs):
            text += para.text + "\n"
            processed += 1
            
            # Update progress periodically
            if i == 0 or processed == total_elements or processed % 20 == 0:
                # Create a new event loop for this thread if needed
                try:
                    loop = asyncio.get_event_loop()
                except RuntimeError:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                
                # Wait for any previous update to complete
                if progress_update_task is not None:
                    if not progress_update_task.done():
                        loop.run_until_complete(progress_update_task)
                
                # Schedule the new progress update
                progress_update_task = loop.create_task(update_progress(processed, total_elements))
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
            text += "\n"
            processed += 1
            
    except Exception as e:
        raise Exception(f"DOCX extraction error: {str(e)}")
    
    return text

async def get_processing_status(job_id: str) -> Optional[Dict[str, Any]]:
    """Get the current processing status for a job"""
    job_data = await get_job(job_id)
    if job_data:
        # Convert to dict for backward compatibility
        return job_data.dict()
    return None

"""
Claude, implement the core document parsing engine for FormMonkey's legal document intelligence system.

Document Processing Pipeline:
- Multi-format parsing: PDF (text/image), DOCX, images, scanned documents
- Intelligent OCR with legal document optimization (Tesseract, cloud OCR APIs)
- Document structure analysis: headers, sections, tables, signature blocks, form fields
- Text extraction with layout preservation and coordinate mapping
- Content classification by document type (contract, form, agreement, etc.)

Dependencies & Integration:
- Import config.py for OCR model paths, processing limits, and quality settings
- Use parser/ocr_utils.py for OCR processing and image enhancement
- Call parser/parser_service.py for document structure analysis
- Import shared/types.ts for ParsedField and FieldType definitions
- Used by routers/upload.py.process_document() for initial document analysis
- Called by routers/parse.py.get_parsed_fields() for field extraction
- Integrates with services/ai_assistance.py for semantic field enhancement
- Calls storage/storage_service.py for temporary file management during processing

Semantic Field Detection:
- Legal field taxonomy: parties, dates, amounts, terms, clauses, obligations
- Pattern recognition for common legal constructs (whereas clauses, definitions, schedules)
- Form field identification: checkboxes, text inputs, dropdowns, signature areas
- Table parsing with column/row semantic understanding
- Cross-reference detection (defined terms, section references, exhibits)

Advanced Document Understanding:
- Legal clause classification and standardization
- Jurisdiction and law type detection (state, federal, international)
- Document hierarchy mapping (main agreement, schedules, amendments)
- Template recognition and field mapping for known document types
- Content validation against legal document standards

AI Integration:
- Modular AI model integration for semantic enhancement
- Local model support for privacy-sensitive documents
- Confidence scoring for all extracted elements
- Alternative interpretation suggestions for ambiguous content
- Context-aware field relationship mapping

Architecture:
- Plugin-based parsing system for different document types
- Scalable processing with job queuing and progress tracking
- Caching layer for repeated document structure patterns
- Error handling with detailed diagnostic information
- Performance optimization with incremental processing capabilities

Output Format:
- Structured document representation with semantic annotations
- Field extraction results with confidence scores and bounding boxes
- Document metadata with classification and quality metrics
- Error reporting with specific location context and suggested corrections
"""

# TODO [0]: Accept uploaded file, perform OCR if needed
# TODO [0.1]: Add comprehensive error handling for malformed documents
# TODO [0.2]: Implement document format detection and validation
# TODO [1]: Detect semantic fields, clauses, tables
# TODO [1.1]: Add confidence scoring for field detection
# TODO [1.2]: Implement fallback detection methods for reliability
# TODO [2]: Return normalized structured field list
# TODO [2.1]: Add regex patterns for common legal document formats
# TODO [2.2]: Implement fuzzy matching for field variations
# TODO [3]: Interface with shared.types and storage utils
# TODO [3.1]: Add caching layer for processed document segments
# TODO [3.2]: Implement parallel processing for large documents
